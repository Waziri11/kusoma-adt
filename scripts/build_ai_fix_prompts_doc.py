#!/usr/bin/env python3
"""Build the Kusoma ADT AI remediation prompt guide."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Kusoma_ADT_AI_Fix_Prompts.docx"


GLOBAL_CONTEXT = """You are editing the self-contained Kusoma Accessible Digital Textbook (ADT) bundle. The source language is sw-TZ. Before changing anything, inspect the target page HTML, content/i18n/sw-TZ/texts.json, content/i18n/sw-TZ/audios.json, the referenced images, and the corresponding original page render or source PDF. Preserve stable data-id values and reading order. When text changes, keep the HTML fallback text and texts.json synchronized. When spoken content changes, regenerate the mapped MP3 with the Microsoft Edge TTS voice sw-TZ-RehemaNeural. Do not edit assets/base.bundle.min.js, assets/base.bundle.local.js, or compiled runtime modules. For image/layout fixes, use object-contain and responsive constraints so the complete source image remains visible. For interactive fill-in-the-blank activities, create real keyboard-focusable inputs tied to data-activity-item and window.correctAnswers; visible underscores alone are not acceptable. After content changes, regenerate assets/offline-preloader.js and increment assets/config.json bundleVersion plus the HTML asset query strings to avoid stale live content. Validate JSON, confirm every data-id still resolves, test read-aloud, keyboard interaction, responsive layout, and GitHub Pages output."""


PROMPTS = [
    ("1. Certification page: complete and accurate read-aloud", "Pages 1–2 / certification and publication metadata", """Inspect the certification/publication page against the original PDF. Correct the fallback HTML, texts.json values, and audio so the read-aloud accurately includes the certificate details, ISBN, approval date, and year. Do not infer missing metadata: transcribe it from the source page. Regenerate each affected MP3 with sw-TZ-RehemaNeural and verify the spoken order matches the visual reading order.""", ["ISBN, date, and year match the source exactly.", "No metadata is skipped or read in English.", "HTML, texts.json, audios.json, and MP3 files remain synchronized."]),
    ("2. Table of contents: pronounce letter sounds correctly", "pg003_sec001.html", """Audit the Table of Contents read-aloud, especially chapter descriptions that list individual letters and consonant clusters. Regenerate audio so letters are spoken as the intended Kiswahili phonemes rather than English letter names. Treat sh, th, mb, ny, ng, nd, and kw as Kiswahili digraphs/clusters and preserve the printed text.""", ["Every listed letter/cluster is pronounced in standard Kiswahili.", "No English alphabet-name pronunciation remains."]),
    ("3. Acknowledgements voice consistency and abbreviations", "pg005_sec001.html", """Replace the Shukurani audio voice with the same voice used for Utangulizi: sw-TZ-RehemaNeural. Correct the pronunciation of the abbreviation TET as separate Kiswahili letter sounds and review KKK or any other all-capital abbreviations against the printed source. Keep the visible text unchanged unless it differs from the PDF.""", ["Shukurani and Utangulizi use the same RehemaNeural voice.", "TET and other abbreviations are clearly pronounced in Kiswahili.", "Names and institutional titles are not anglicized."]),
    ("4. Restore complete images on affected pages", "pg008, pg033, pg071, and pg080", """Compare every image on pages 8, 33, 71, and 80 with the original page renders. Fix the HTML layout so no image is clipped by fixed heights, oversized flex rows, object-cover, or overflow-hidden containers. Use responsive grid/flex sizing, max-width:100%, height:auto, and object-contain. On pg080 specifically, the complete shoka, tai, and ndoo must fit inside the page at the same time; neither the axe head/handle nor the bucket edge may be cut off.""", ["All source objects are visible at desktop and mobile widths.", "Images retain aspect ratio and do not overlap text or navigation.", "No unnecessary image file recropping is used when CSS sizing is the cause."]),
    ("5. Page 54 image descriptions and missing images", "pg054_sec001.html", """Compare pg054 with the original PDF and inventory every visual item. Put image descriptions in the same logical reading order as the images and add any source image omitted from the ADT page. Each image must have a unique data-id, an accurate Kiswahili description in texts.json, an audio mapping, and a RehemaNeural MP3. Do not merge descriptions for separate images.""", ["Every source image appears once.", "Screen-reader and read-aloud order matches the visual order.", "Descriptions identify the actual subject without speculation."]),
    ("6. Match image-number typography", "pg010_sec001.html compared with pg009_sec001.html", """Update the visible numbers that label the images on pg010 so their font family, weight, size, color, alignment, and placement match the equivalent image labels on pg009. Preserve the underlying number data-ids and activity logic. Confirm the labels remain readable and do not cover important image content at small widths.""", ["pg010 image labels visually match pg009.", "No data-id or navigation behavior changes."]),
    ("7. Speak all numerals and exercise numbers in Kiswahili", "All pages; priority pg014, pg015, pg062 and every ‘Zoezi la …’", """Find every audio-bearing text that contains a numeral, including leading question numbers, phrases such as ‘Zoezi la 3’, and image references such as ‘picha namba 1, 2, 3 na 4’. Keep numerals visible in the book, but synthesize spoken text using Kiswahili number words and appropriate ordinals: 1→moja/kwanza, 2→mbili/pili, 3→tatu, 4→nne, etc., according to grammar. Regenerate normal and easy-read MP3s with sw-TZ-RehemaNeural. On pg062, the numbered instructions must say ‘moja’, ‘mbili’, and ‘tatu’, never ‘one’, ‘two’, or ‘three’.""", ["No English number words occur in read-aloud.", "Every ‘Zoezi la …’ number is pronounced naturally in Kiswahili.", "Normal and mapped easy-read audio are both updated."]),
    ("8. Correct vowel and consonant sound pronunciation", "pg016, pg017, pg024, and pg067", """Regenerate educational phonics audio so isolated vowels and consonants are pronounced as Kiswahili sounds, not English alphabet names. On pg016, isolated ‘a’ and ‘u’ must sound /a/ and /u/ (never ‘yu’). On pg017 and pg024, examples such as b–bu, m–mu, k–ka, d–di, and n–ni must pause clearly between the consonant sound and syllable. On pg067, pronounce the consonant clusters sh, th, mb, ny, ng, nd, and kw as the intended Kiswahili sounds.""", ["Letters are phonemes, not English letter names.", "Cluster audio remains intelligible at normal playback speed.", "Printed instructional text is preserved."]),
    ("9. Correct duplicated word in question 6", "pg021_sec001.html, pg021_n0010", """Compare the question-6 word list with the PDF. The current text ‘bomoa boma’ must be corrected to the exact two intended words from the source; do not leave a duplicated ‘bomoa’ or silently guess. Update the inline HTML, texts.json, any easy-read entry, and the mapped audio. Preserve the sentence-writing activity and its input.""", ["The two words match the source page.", "Text and audio agree exactly."]),
    ("10. Correct picture name under the letter d section", "pg023_sec001.html", """Audit the pictures in the ‘Sauti ya herufi d’ section against the source. Correct the wrongly spoken picture name so the cup is identified as ‘Kikombe’ where required by the source, and ensure each image’s alt/description, texts.json entry, and MP3 refer to the same object. Do not rename unrelated d-word pictures without source evidence.""", ["The cup is called ‘Kikombe’ in visual accessibility text and audio.", "Each image description matches its own file."]),
    ("11. Repair the two cut coconut pieces", "pg026_sec001.html and images/pg026_im001.jpg", """Inspect pg026 against the PDF. The first visual must show one whole coconut and two complete cut coconut pieces. If the source asset already contains the complete composition, fix CSS/object fitting; if the extracted asset is genuinely incomplete, replace it with a correctly extracted crop from the original page render. Update the Kiswahili alt description to ‘Nazi nzima na vipande viwili vya nazi vilivyokatwa’ and retain the existing stable image data-id.""", ["Both cut pieces are fully visible.", "No piece is duplicated, stretched, or clipped.", "The description and image agree."]),
    ("12. Correct page 28 picture identity and make blanks interactive", "pg028_sec001.html", """Compare every page-28 picture with the PDF and correct the first picture’s name/description to the actual object shown. Convert every visible underscore group into a real fillable input while preserving the fixed letters and intended completed words. Use keyboard-focusable text inputs with unique data-activity-item values and correct window.correctAnswers entries. Keep separate sr-only read-aloud strings that say ‘dashi’ for each blank. Regenerate the blank-sentence MP3s with sw-TZ-RehemaNeural; the audio must announce each blank as ‘dashi’, not read underscores or omit the blank.""", ["All visible blanks can be typed into and submitted.", "The first image name matches the source.", "Read-aloud says ‘dashi’ at every blank.", "Correct answers validate the intended words."]),
    ("13. Make question 8 dumu and kuku blanks interactive", "pg027_sec001.html, question 8", """Replace the decorative underline spans in question 8 with real fill-in-the-blank controls. For dumu, show blank–u–blank–u and accept d then m. For kuku, show k–blank–k–blank and accept u then u. Bind the inputs to the existing activity using data-activity-item and correctAnswers. Keep pg027_n0100 and pg027_n0102 as screen-reader/read-aloud text using the word ‘dashi’, and regenerate those MP3s with sw-TZ-RehemaNeural.""", ["Exactly four editable blanks appear.", "Answers d, m, u, u validate correctly.", "Tab order follows dumu then kuku.", "Audio says ‘dashi’ for all four blanks."]),
    ("14. Correct pronunciation of ‘nile’", "pg032_sec001.html", """Locate every normal and easy-read occurrence of ‘nile’ on pg032. Regenerate only the affected MP3s with sw-TZ-RehemaNeural, using natural Kiswahili pronunciation and sentence context. Do not respell the visible word merely to manipulate TTS unless the source text itself is wrong.""", ["‘nile’ is clearly distinguishable and naturally stressed.", "Visible spelling remains source-accurate."]),
    ("15. Correct sentence pronunciation: Debe limejaa gololi", "pg043_sec001.html, question 8", """Locate the sentence ‘Debe limejaa gololi’ in question 8 and regenerate its normal and easy-read audio with sw-TZ-RehemaNeural. Preserve the exact visible sentence and ensure the voice clearly pronounces ‘Debe’, ‘limejaa’, and ‘gololi’ as three Kiswahili words without merging, skipping, or anglicizing them.""", ["The complete sentence is spoken once in the correct order.", "No word is dropped or substituted."]),
    ("16. Audit syllable and word-segmentation audio", "pg053, pg056, pg070, pg078, and pg081", """Regenerate phonics audio with explicit pauses that reflect the pedagogy. On pg053, pronounce each sound that forms the shown syllables/words. On pg056, correct the letter sounds in Zoezi la Kwanza. On pg070, include every example syllable, especially ‘ni’, and pronounce the third example for ‘thelathini’ as th–e–l–a–th–i–n–i. On pg078, preserve consonant clusters: ngumi must be segmented as ngu + mi for syllables and ng + u + m + i only when the instruction explicitly asks for letter sounds. On pg081, do not over-segment syllables: ndama must be nda + ma, not nda + m + a, when the task asks for syllables.""", ["Segmentation follows the wording of each instruction.", "No example syllable is omitted.", "Clusters such as ng and nd stay intact when functioning as one sound."]),
    ("17. Correct difficult Kiswahili word pronunciations", "pg061, pg062, pg085, and pg086", """Audit all normal and easy-read MP3s on these pages for standard Kiswahili pronunciation. Prioritize ‘mbalimbali’, ‘mfanyabiashara’, and ‘kuving’arisha’. Use the printed apostrophe in ng’ consistently, preserve complete words, and regenerate affected audio with sw-TZ-RehemaNeural. Do not replace correct Kiswahili spellings with English-friendly phonetic spellings in texts.json.""", ["Target words are complete and intelligible.", "ng’ is pronounced as the intended Kiswahili sound.", "No English voice or accent remains."]),
    ("18. Correct picture names on the kw page", "pg082_sec001.html", """Compare all pg082 pictures with the PDF. Verify the two hoof images and the pictures currently described as Kware, Bunduki, and Kwokoo. Replace any incorrect name with the source-accurate Kiswahili noun, then synchronize image alt text, texts.json, easy-read content if present, audio mappings, and RehemaNeural MP3s. Do not preserve a name solely because it begins with kw if the picture shows a different object.""", ["Every image name is visually defensible from the source.", "Descriptions and audio use the same noun."]),
    ("19. Correct the mosquito-scene description", "pg093_sec001.html, pg093_im001", """Replace the inaccurate image description that refers to flies or generic insects. The source scene shows a sleeping boy surrounded by mosquitoes while not protected by a mosquito net. Write a concise Kiswahili description stating that clearly, update the matching texts.json entry, and regenerate the image-description MP3 with sw-TZ-RehemaNeural. Do not add details that are not visible.""", ["The description explicitly says ‘mbu’ and notes the missing protective net.", "Audio and alt/description text match."]),
    ("20. Global Kiswahili language and voice audit", "Entire ADT bundle", """Scan every mapped text/audio pair for English speech, English number names, inconsistent voices, skipped words, and mispronounced Kiswahili phonics. Use sw-TZ-RehemaNeural consistently. Preserve the source language sw-TZ and standard Tanzanian Kiswahili. Do not translate proper names or institutional names, but expand or carefully synthesize abbreviations where needed. Update both normal and mapped easy-read audio, then produce a machine-readable audit report listing each regenerated text-id, source text, spoken synthesis text, voice, and output filename.""", ["All mapped audio uses the approved Swahili voice.", "No English number or alphabet-name pronunciation remains.", "Audit report accounts for every changed MP3."]),
    ("21. Publish without stale GitHub Pages content", "Deployment and cache-busting", """After all content, audio, and layout fixes are complete, regenerate assets/offline-preloader.js from the current JSON files. Increment assets/config.json bundleVersion and update the offline-preloader.js and base.bundle.local.js query-string version in every page HTML, including index.html. Validate JSON, commit all intended files, push main, wait for the GitHub Pages workflow, and compare SHA hashes of representative live files with local files using cache-busting query parameters.""", ["GitHub Pages workflow succeeds.", "Live representative HTML and MP3 hashes match local files.", "The live config and preloader expose the new bundle version.", "A normal refresh loads the new assets after CDN/browser cache expiry; a versioned URL loads them immediately."]),
]


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for style_name, size, color, before, after in (
    ("Title", 24, "17365D", 0, 8),
    ("Heading 1", 16, "17365D", 14, 7),
    ("Heading 2", 12, "2E74B5", 10, 5),
):
    style = styles[style_name]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Kusoma ADT — AI Remediation Prompts")
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Descriptive implementation prompts derived from KUSOMA KUNDIRIK.doc and verified against the ADT bundle")
run.italic = True
run.font.color.rgb = RGBColor(89, 89, 89)

doc.add_heading("How to use this guide", level=1)
doc.add_paragraph("Copy the Global Context once, then append the relevant numbered prompt. Give the AI access to the complete ADT bundle and the original PDF/page renders. Each prompt is written as an implementation task, not merely a diagnosis.")

doc.add_heading("Global Context — include with every prompt", level=1)
table = doc.add_table(rows=1, cols=1)
table.autofit = False
table.columns[0].width = Inches(6.8)
cell = table.cell(0, 0)
shade(cell, "EAF2F8")
set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
cell.paragraphs[0].add_run(GLOBAL_CONTEXT)

doc.add_heading("Page-specific and global prompts", level=1)

for title_text, scope, prompt, checks in PROMPTS:
    doc.add_heading(title_text, level=2)
    p = doc.add_paragraph()
    label = p.add_run("Scope: ")
    label.bold = True
    p.add_run(scope)

    p = doc.add_paragraph()
    label = p.add_run("Prompt: ")
    label.bold = True
    p.add_run(prompt)

    p = doc.add_paragraph()
    label = p.add_run("Acceptance checks")
    label.bold = True
    for check in checks:
        doc.add_paragraph(check, style="List Bullet")

doc.add_heading("Required final verification", level=1)
for item in (
    "Open each changed page at desktop and narrow/mobile widths; verify no clipping, overlap, or missing controls.",
    "Run keyboard-only interaction through every changed activity and submit correct and incorrect answers.",
    "Listen to every regenerated MP3 in context and verify the approved RehemaNeural voice and standard Kiswahili pronunciation.",
    "Validate assets/config.json, content/pages.json, texts.json, audios.json, glossary.json, and videos.json as valid JSON.",
    "Regenerate offline-preloader.js, increment bundleVersion, publish, and compare representative live/local hashes.",
):
    doc.add_paragraph(item, style="List Bullet")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run("Kusoma ADT AI remediation prompt guide")
footer_run.font.size = Pt(8)
footer_run.font.color.rgb = RGBColor(117, 117, 117)

doc.save(OUTPUT)
print(OUTPUT)
